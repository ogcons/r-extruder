from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework import status, generics
from .models import RScript
from .serializers import RScriptSerializer
from django.conf import settings
import os
import rpy2.robjects as robjects
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches

class RScriptListCreateView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        try:
            serializer = RScriptSerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class RScriptRetrieveView(APIView):
    def get(self, request, pk, *args, **kwargs):
        try:
            r_script = RScript.objects.get(pk=pk)
        except RScript.DoesNotExist:
            return Response({"error": "RScript not found"}, status=status.HTTP_404_NOT_FOUND)

        serializer = RScriptSerializer(r_script)
        return Response(serializer.data, status=status.HTTP_200_OK)

class RScriptListCreateView(generics.ListCreateAPIView):
    queryset = RScript.objects.all()
    serializer_class = RScriptSerializer

class RunRScriptView(APIView):
    def post(self, request, *args, **kwargs):
        BASE_DIR = settings.BASE_DIR

        try:
            # Get the RScript instance from the request data
            r_script_serializer = RScriptSerializer(data=request.data)
            if r_script_serializer.is_valid():
                r_script = r_script_serializer.save()
            else:
                return Response({"error": "Invalid RScript data"}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Ensure the 'media' directory exists
        media_directory = os.path.join(BASE_DIR, 'media')
        if not os.path.exists(media_directory):
            os.makedirs(media_directory)

        # Use absolute paths for the R script and plot
        r_script_path = os.path.join(BASE_DIR, 'media', str(r_script.script))
        plot_file_path = os.path.join(BASE_DIR, 'media', 'plot.png')

        try:
            # Initialize R environment
            r = robjects.r

            # Read the content of the R script
            with open(r_script_path, 'r') as file:
                r_script_code = file.read()

            # Check if the R script contains the plot function
            if 'plot(' in r_script_code:
                # Add code for saving the plot as PNG before the plot function
                r_script_code = r_script_code.replace('plot(', 'png(\'media/plot.png\')\nplot(')

                # Add code for dev.off() after the plot function
                r_script_code += "\ndev.off()"

            # Execute modified R script
            r(r_script_code)

            # You can return a success message or any other relevant information
            return Response({"message": "R script executed successfully", "plot_url": "media/plot.png"})
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def test_view(request):
    try:
        print("Inside generate_word_document view")
        # Path to the plot.png file in the media directory
        plot_file_path = os.path.join(settings.BASE_DIR, 'media', 'plot.png')

        # Create a new Word document
        document = Document()

        # Add a paragraph with some text
        content = "Hello, this is the content of the Word document."
        document.add_paragraph(content)

        # Add an image to the Word document
        document.add_picture(plot_file_path, width=Inches(4))  # Adjust the width as needed

        # Create a response with the appropriate content type for Word documents
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        # Set the content-disposition header to prompt the user to download the Word document
        response['Content-Disposition'] = 'attachment; filename=document_with_image.docx'

        # Save the Word document to the response
        document.save(response)

        # Remove the plot.png file
        os.remove(plot_file_path)

        return response
    except Exception as e:
        return HttpResponse(f"Error: {str(e)}", status=500)
